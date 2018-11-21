# _*_ encoding:utf-8 _*_
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from docx.shared import Inches

time_now = datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")
time_now = time_now.decode('utf-8')


def docx(now_time=time_now, start_time=time_now, end_time=time_now, jiedian="Al", user=u"管理员", datas=None):
    document = Document()
    document.add_heading(u'电能质量报表', 0)
    document.styles['Normal'].font.name = u'黑体'
    p = document.add_paragraph()
    # p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(u'打印时间:' + now_time + "\n")
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(11)
    run = p.add_run(u'统计时段:' + start_time + "--" + end_time + "\n")
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(11)
    run = p.add_run(u'节点名称:' + jiedian + "\n")
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(11)
    run = p.add_run(u'报表制作:' + user + "\n")
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(11)
    run = p.add_run(u"依据电能质量的国家标准，该节点的电能质量综合评估参数如下：" + "\n")
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(11)

    # pic = document.add_picture('logo1.PNG')
    # last_paragraph = document.paragraphs[-1]
    # last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置

    rows = 2
    cols = len(datas)
    table = document.add_table(rows=rows, cols=cols + 1, style="Table Grid")  # 添加1行10列的表格

    for i in range(rows):
        tr = table.rows[i]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), "250")
        trPr.append(trHeight)  # 表格高度设置
    # table.autofit = False
    col = table.columns[1]
    col.width = Inches(5)
    # arr = [u'', u"优", u"良上", u"良好", u"良下", u"一般", u"较差", u"差", u"很差", u"极差"]
    heading_cells = table.rows[0].cells
    arr_value = table.rows[1].cells
    for i in range(cols):
        p = heading_cells[i + 1].paragraphs[0]
        p2 = arr_value[i + 1].paragraphs[0]
        run = p.add_run(datas[i].keys())
        run2 = p2.add_run(str(datas[i].values()[0]))
        run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
        run.font.size = Pt(11)  # 字体大小设置，和word里面的字号相对应
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run2.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
        run2.font.size = Pt(11)  # 字体大小设置，和word里面的字号相对应
        p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.cell(1, 0).text = u'每小时隶属平均最大最小值'
    table.cell(0, 0).text = u'时间'


    document.save('datas.docx')


def read_file(file_name, size):
    with open(file_name, mode='rb') as fp:
        while True:
            c = fp.read(size)
            if c:
                yield c
            else:
                break

# docx(datas=[{"1":"2"},{"2":"3"}])
